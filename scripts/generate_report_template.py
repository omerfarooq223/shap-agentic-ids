import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_ieee_report():
    print("Generating IEEE Final Report Template (.docx)...")
    doc = Document()
    
    # Title
    title = doc.add_heading('SHAP-Explained Agentic Intrusion Detection System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Author
    author = doc.add_paragraph('Muhammad Umar Farooq\nRoll Number: F20233763310\nCourse: AI-374 | Information Security')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph('This report presents an intrusion detection system (IDS) that combines Random Forest for anomaly detection, SHAP for feature attribution, and a LangGraph-powered LLM agent for threat verification and narrative generation. The system addresses the black-box nature of traditional ML-IDS and the hallucination risks of LLM-only approaches.')
    
    # Sections
    sections = [
        "1. Introduction",
        "2. Literature Review & Gap Analysis",
        "3. System Architecture",
        "4. Methodology",
        "  4.1 Dataset & Preprocessing",
        "  4.2 Handling Class Imbalance with SMOTE",
        "  4.3 Model Training (Random Forest)",
        "  4.4 Explainability Layer (SHAP)",
        "  4.5 Agentic Reasoning (LangGraph + GROQ)",
        "5. Experimental Setup & Results",
        "6. Discussion & Limitations",
        "7. Conclusion",
        "8. References"
    ]
    
    for section in sections:
        level = 1 if not section.startswith(" ") else 2
        doc.add_heading(section.strip(), level=level)
        doc.add_paragraph('[Insert content here based on your system design and proposal...]')
    
    filename = 'IS_Project_Final_Report.docx'
    doc.save(filename)
    print(f"Report template successfully saved to: {os.path.abspath(filename)}")
    print("You can now open this file in Microsoft Word and fill in the specifics.")

if __name__ == "__main__":
    generate_ieee_report()
